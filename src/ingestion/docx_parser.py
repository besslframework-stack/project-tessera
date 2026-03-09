"""Parse DOCX files into LlamaIndex Documents."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from llama_index.core.schema import Document

logger = logging.getLogger(__name__)


def parse_docx_file(file_path: Path) -> list["Document"]:
    """Parse a DOCX file into Documents, splitting by headings.

    Each heading creates a new section. Paragraphs under a heading
    are grouped into one Document per section.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.warning("python-docx not installed, skipping DOCX: %s", file_path)
        return []

    from llama_index.core.schema import Document

    try:
        doc = DocxDocument(str(file_path))
    except Exception as exc:
        logger.error("Failed to open DOCX %s: %s", file_path, exc)
        return []

    documents: list[Document] = []
    current_section = "Introduction"
    current_text: list[str] = []
    section_index = 0

    def _flush() -> None:
        nonlocal section_index
        text = "\n".join(current_text).strip()
        if not text:
            return
        documents.append(
            Document(
                text=text,
                metadata={
                    "source_path": str(file_path),
                    "file_name": file_path.name,
                    "file_type": "docx",
                    "section": current_section,
                    "section_index": section_index,
                },
            )
        )
        section_index += 1

    for para in doc.paragraphs:
        # Check if paragraph is a heading
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            _flush()
            current_section = para.text.strip() or "Untitled Section"
            current_text = []
        else:
            text = para.text.strip()
            if text:
                current_text.append(text)

    _flush()
    return documents
